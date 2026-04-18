"""Generación de borrador Word — Brief 3.8, tarea S4.2 (herramienta determinista, sin LLM)."""

from __future__ import annotations

import io
import json
import re
from datetime import date, datetime
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table

from apps.agents.models.caso import Caso
from apps.agents.models.fase_output import FaseOutput, WordGeneratorInput, WordOutput
from apps.agents.models.hecho import Hecho

FASES_DOCUMENTO: tuple[tuple[str, str], ...] = (
    ("0E", "Análisis ético (0E)"),
    ("0A", "Auditoría de hechos (0A)"),
    ("0C", "Estrategia inicial (0C)"),
    ("1A", "Preguntas críticas al cliente (1A)"),
    ("1C", "Teoría del caso (1C)"),
    ("2A", "Motor probabilístico (2A)"),
    ("5A", "Simulación adversarial (5A)"),
    ("GEN", "Generación / borrador (GEN)"),
)

_VERIFICAR_SPLIT = "VERIFICAR"


def _margins_cm(section: Any, cm: float = 2.5) -> None:
    section.top_margin = Cm(cm)
    section.bottom_margin = Cm(cm)
    section.left_margin = Cm(cm)
    section.right_margin = Cm(cm)


def _set_run_font_times(run: Any, *, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    r = run._element.rPr
    if r is not None:
        rFonts = r.rFonts
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            r.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")


def _aplicar_estilo_base(doc: Document) -> None:
    for sec in doc.sections:
        _margins_cm(sec, 2.5)
    estilo = doc.styles["Normal"]
    estilo.font.name = "Times New Roman"
    estilo.font.size = Pt(12)
    pf = estilo.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5


def _safe_nombre_archivo(nombre: str) -> str:
    limpio = re.sub(r'[<>:"/\\|?*]', "_", nombre.strip())
    limpio = re.sub(r"\s+", "_", limpio)
    return (limpio[:120] or "caso").strip("_") or "caso"


def _nombre_archivo_borrador(metadata: Caso) -> str:
    fecha = date.today().isoformat()
    base = _safe_nombre_archivo(metadata.nombre_caso)
    return f"borrador_{base}_{fecha}.docx"


def _resolver_fase(outputs_fases: dict[str, FaseOutput], codigo: str) -> FaseOutput | None:
    if codigo in outputs_fases:
        return outputs_fases[codigo]
    clave_lower = codigo.lower()
    for k, fo in outputs_fases.items():
        if str(k).lower() == clave_lower:
            return fo
        if str(fo.fase).upper() == codigo.upper():
            return fo
    return None


def _contenido_a_texto(contenido: dict[str, Any]) -> str:
    try:
        return json.dumps(contenido, ensure_ascii=False, indent=2)
    except (TypeError, ValueError):
        return str(contenido)


def _add_paragraph_with_verificar(doc: Document, texto: str) -> None:
    """Inserta párrafo; la subcadena VERIFICAR aparece en rojo."""
    if not texto.strip():
        doc.add_paragraph()
        return
    for bloque in texto.split("\n"):
        linea = bloque if bloque else " "
        p = doc.add_paragraph()
        partes = linea.split(_VERIFICAR_SPLIT)
        for i, parte in enumerate(partes):
            if parte:
                r = p.add_run(parte)
                _set_run_font_times(r, bold=False)
            if i < len(partes) - 1:
                r_v = p.add_run(_VERIFICAR_SPLIT)
                _set_run_font_times(r_v, bold=False)
                r_v.font.color.rgb = RGBColor(220, 20, 20)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5


def _add_titulo_seccion(doc: Document, numero: int, titulo: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{numero}. {titulo}")
    run.bold = True
    _set_run_font_times(run, bold=True)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5


def _tiene_cifras_liquidacion(textos: list[str]) -> bool:
    patron = re.compile(
        r"(?:\bCOP\b|\bUSD\b|\$|€|\b\d{1,3}(?:\.\d{3})+(?:,\d+)?\b|\b\d+,\d{2}\b|\b\d{4,}\b)",
        re.IGNORECASE,
    )
    return any(patron.search(t) for t in textos if t)


def _add_tabla_liquidacion_si_aplica(
    doc: Document,
    numero_seccion: int,
    hechos: list[Hecho],
    textos_fases: list[str],
) -> bool:
    """Devuelve True si insertó tabla."""
    bloques: list[str] = [h.contenido for h in hechos]
    if not _tiene_cifras_liquidacion(bloques + textos_fases):
        return False
    _add_titulo_seccion(doc, numero_seccion, "Referencias numéricas y liquidación (cuantía)")
    p = doc.add_paragraph()
    r = p.add_run("Tabla de referencia: contrastar con prueba y anexos originales.")
    r.italic = True
    _set_run_font_times(r, bold=False)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()

    filas: list[tuple[str, str]] = []
    for h in hechos:
        if re.search(r"\d", h.contenido):
            filas.append((f"Hecho ({h.fase_origen})", h.contenido[:500]))
    if not filas:
        filas.append(("Nota", "Se detectaron cifras en fases; detalle en secciones correspondientes."))

    nfil = min(len(filas), 20)
    tabla: Table = doc.add_table(rows=1 + nfil, cols=2)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tabla.rows[0].cells
    hdr[0].text = "Concepto"
    hdr[1].text = "Detalle en expediente"
    for i in range(nfil):
        c0, c1 = tabla.rows[i + 1].cells
        c0.text = filas[i][0]
        c1.text = filas[i][1]
    doc.add_paragraph()
    return True


def _add_tabla_protocolo_verificacion(doc: Document, numero_seccion: int) -> None:
    _add_titulo_seccion(doc, numero_seccion, "Protocolo de verificación (cierre)")
    doc.add_paragraph()
    items: list[tuple[str, str]] = [
        ("Coherencia interna de hechos", "Revisar sección de hechos y pretensiones."),
        ("Sustento fáctico de pretensiones", "Contrastar cada pretensión con anexos citados."),
        ("Citas normativas pertinentes", "VERIFICAR vigencia y concordancia con tipo de acción."),
        ("Identificación de partes", "VERIFICAR documentos de identificación y poderes."),
        ("Medidas cautelares", "VERIFICAR requisitos legales y competencia."),
        ("Números y cifras", "VERIFICAR aritmética y unidades frente a prueba."),
    ]
    tabla: Table = doc.add_table(rows=1 + len(items), cols=2)
    tabla.style = "Table Grid"
    h = tabla.rows[0].cells
    h[0].text = "Ítem de verificación"
    h[1].text = "Estado / acción sugerida"
    for i, (a, b) in enumerate(items):
        row = tabla.rows[i + 1].cells
        row[0].text = a
        row[1].text = b


def _estimar_paginas(texto_total: int) -> int:
    if texto_total <= 0:
        return 1
    return max(1, min(200, texto_total // 2200 + 1))


def generate_word_document(inp: WordGeneratorInput) -> WordOutput:
    """
    Construye un .docx determinista a partir de outputs de fase y hechos.

    No llama a modelos de lenguaje.
    """
    advertencias: list[str] = []
    doc = Document()
    _aplicar_estilo_base(doc)

    meta = inp.metadata_caso
    titulo = doc.add_paragraph()
    tr = titulo.add_run(f"Borrador jurídico — {meta.nombre_caso}")
    tr.bold = True
    _set_run_font_times(tr, bold=True)
    titulo.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    titulo.paragraph_format.line_spacing = 1.5

    sub = doc.add_paragraph()
    sr = sub.add_run(
        f"Caso ID: {inp.caso_id} · Tipo de acción: {meta.tipo_accion.value} · "
        f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')} (herramienta determinista)"
    )
    _set_run_font_times(sr, bold=False)
    sub.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    sub.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()

    textos_para_liquidacion: list[str] = []
    caracteres_totales = 0

    for idx, (codigo, etiqueta) in enumerate(FASES_DOCUMENTO, start=1):
        _add_titulo_seccion(doc, idx, etiqueta)
        fo = _resolver_fase(inp.outputs_fases, codigo)
        if fo is None:
            _add_paragraph_with_verificar(doc, "[PENDIENTE - fase no completada]")
            advertencias.append(f"Fase {codigo} sin output en outputs_fases.")
            continue
        if not fo.aprobado_abogado:
            _add_paragraph_with_verificar(doc, "[PENDIENTE - fase no completada]")
            advertencias.append(
                f"Fase {codigo}: hay output pero sin aprobación del abogado; no se incorpora el contenido al borrador."
            )
            continue
        bloque = _contenido_a_texto(fo.contenido)
        caracteres_totales += len(bloque)
        textos_para_liquidacion.append(bloque)
        _add_paragraph_with_verificar(doc, bloque)
        doc.add_paragraph()

    numero_seccion = len(FASES_DOCUMENTO) + 1
    _add_titulo_seccion(doc, numero_seccion, "Hechos del caso (resumen)")
    if not inp.hechos:
        _add_paragraph_with_verificar(doc, "[PENDIENTE - fase no completada]")
        advertencias.append("No se recibieron hechos en WordGeneratorInput.")
    else:
        for h in inp.hechos:
            linea = (
                f"- [{h.fase_origen}] ({h.estatus_epistemico.value}) "
                f"{h.contenido[:2000]}"
                + ("…" if len(h.contenido) > 2000 else "")
            )
            caracteres_totales += len(linea)
            _add_paragraph_with_verificar(doc, linea)
    doc.add_paragraph()

    numero_seccion += 1
    if _add_tabla_liquidacion_si_aplica(doc, numero_seccion, inp.hechos, textos_para_liquidacion):
        advertencias.append(
            "Se insertó tabla de referencias numéricas: VERIFICAR cifras contra prueba original."
        )
        numero_seccion += 1
    else:
        doc.add_paragraph()

    _add_tabla_protocolo_verificacion(doc, numero_seccion)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    archivo = buf.read()

    paginas = _estimar_paginas(caracteres_totales + len(archivo) // 80)
    nombre = _nombre_archivo_borrador(meta)

    return WordOutput(
        archivo_bytes=archivo,
        nombre_archivo=nombre,
        paginas_estimadas=paginas,
        advertencias=advertencias,
    )
